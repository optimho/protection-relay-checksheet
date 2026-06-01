#!/usr/bin/env python3
"""
Generate ABB Relion REG670 2.0 (ANSI) Generator Protection IED test documents:
  1. REG670_Test_Check_Sheet.docx       - maintenance check sheet
  2. REG670_Test_Results_Template.xlsx  - live results template with auto pass/fail

Follows the structural convention of the existing SEL-300G pair in this directory,
adapted for ABB/Relion terminology (PCM600, LHMI, INTERRSIG self-supervision, etc.).
"""

import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.formatting.rule import CellIsRule
from openpyxl.workbook.defined_name import DefinedName
from openpyxl.utils import get_column_letter

HERE = os.path.dirname(os.path.abspath(__file__))

# ---- Brand palette ---------------------------------------------------------
BRAND = "1F4E79"      # deep blue - headers
ACCENT = "D9E2F3"     # light blue - label cells
NOTE_BG = "FFF2CC"    # yellow - callouts / tolerance
PASS_FILL = "C6EFCE"; PASS_TXT = "006100"
FAIL_FILL = "FFC7CE"; FAIL_TXT = "9C0006"
WHITE = "FFFFFF"; BLACK = "000000"; INPUT_BLUE = "0000FF"

# ===========================================================================
#  WORD CHECK SHEET
# ===========================================================================

def set_cell_shading(cell, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def style_cell_text(cell, bold=False, color=None, size=11, align=None):
    for p in cell.paragraphs:
        if align is not None:
            p.alignment = align
        for r in p.runs:
            r.font.bold = bold
            r.font.size = Pt(size)
            r.font.name = "Calibri"
            if color:
                r.font.color.rgb = RGBColor.from_string(color)
        if not p.runs and (bold or color):
            r = p.add_run("")
            r.font.bold = bold


def add_para(doc, text, style=None, bold=False, color=None, size=11, italic=False,
             align=None, space_after=4):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_note(doc, text):
    """Yellow callout note box (single-cell shaded table)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, NOTE_BG)
    p = cell.paragraphs[0]
    r = p.add_run("NOTE: " + text)
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.font.italic = True
    set_table_borders(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def add_checks(doc, items):
    for it in items:
        add_para(doc, "☐  " + it, size=11, space_after=2)


def set_table_borders(tbl, color="BFBFBF", sz="4"):
    tblPr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def make_table(doc, headers, rows, col_widths=None, label_first_col=False,
               header_fill=BRAND, fixed_data_rows=0):
    """Generic table. `rows` is list-of-lists. If fixed_data_rows>0, add that many
    blank rows after the supplied rows. label_first_col shades col0 with ACCENT."""
    ncol = len(headers)
    nrows = 1 + len(rows) + fixed_data_rows
    tbl = doc.add_table(rows=nrows, cols=ncol)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    set_table_borders(tbl)
    # header
    hdr = tbl.rows[0]
    set_repeat_header(hdr)
    for j, h in enumerate(headers):
        c = hdr.cells[j]
        c.text = h
        set_cell_shading(c, header_fill)
        style_cell_text(c, bold=True, color=WHITE, size=10,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
    # data rows
    for i, rowdata in enumerate(rows, start=1):
        for j, val in enumerate(rowdata):
            c = tbl.rows[i].cells[j]
            c.text = str(val)
            style_cell_text(c, size=10)
            if label_first_col and j == 0:
                set_cell_shading(c, ACCENT)
                style_cell_text(c, bold=True, size=10)
    if col_widths:
        for row in tbl.rows:
            for j, w in enumerate(col_widths):
                if j < len(row.cells):
                    row.cells[j].width = Cm(w)
    return tbl


def make_label_value_table(doc, pairs, label_w=6.5, val_w=10.5):
    """Two-column label/value table with shaded label column."""
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(tbl)
    for i, (lab, val) in enumerate(pairs):
        lc = tbl.rows[i].cells[0]
        lc.text = lab
        set_cell_shading(lc, ACCENT)
        style_cell_text(lc, bold=True, size=10)
        vc = tbl.rows[i].cells[1]
        vc.text = val
        style_cell_text(vc, size=10)
        lc.width = Cm(label_w)
        vc.width = Cm(val_w)
    return tbl


def build_docx():
    doc = Document()

    # --- page setup: A4 portrait, 1 cm margins ---
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(1.0))

    # base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # heading styles -> brand colour
    for hn, sz in (("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11)):
        st = doc.styles[hn]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BRAND)

    # --- title block ---
    add_para(doc, "PROTECTION RELAY TEST CHECK SHEET", bold=True, color=BRAND,
             size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "ABB Relion REG670 2.0 (ANSI) Generator Protection IED", bold=True,
             size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Routine / Maintenance Test Procedure", italic=True, size=11,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # --- Document Control ---
    doc.add_heading("Document Control", level=2)
    make_table(doc,
               ["Revision", "Date", "Description", "Prepared By", "Approved By"],
               [["1.0", "", "Initial issue", "", ""]],
               col_widths=[2.0, 2.5, 6.0, 3.0, 3.0], fixed_data_rows=2)

    # --- Test Summary & Findings ---
    doc.add_heading("Test Summary & Findings", level=2)
    add_para(doc, "Filled in by the tester at the end of the visit. Captures the "
             "overall outcome and any noteworthy findings.", italic=True, size=10)
    make_label_value_table(doc, [
        ("Overall test result",
         "☐ PASS — IED returned to service    "
         "☐ PASS with defects (see findings)    ☐ FAIL / withdrawn"),
        ("Test completion date", ""),
        ("Summary & key findings", ""),
        ("Self-supervision status on departure",
         "☐ Normal (green) LED steady, no FAIL/WARNING    ☐ Issue noted (see findings)"),
    ])

    # ======================= SECTION 1 ===============================
    doc.add_heading("1. Test Information", level=1)

    doc.add_heading("1.1 Site & Project", level=3)
    make_label_value_table(doc, [
        ("Site / Power station", ""),
        ("Asset identifier (generator / unit)", ""),
        ("Work order / Job number", ""),
        ("Reason for test", "Routine / Maintenance"),
        ("Test date (start)", ""),
        ("Test date (finish)", ""),
    ])

    doc.add_heading("1.2 Personnel", level=3)
    make_label_value_table(doc, [
        ("Tester (name)", ""),
        ("Tester signature", ""),
        ("Witness / 2nd person", ""),
        ("Witness signature", ""),
        ("Authorising engineer", ""),
    ])

    doc.add_heading("1.3 IED Under Test", level=3)
    make_label_value_table(doc, [
        ("IED model", "ABB Relion REG670 2.0 (ANSI)"),
        ("IED tag / functional ID", ""),
        ("Serial number (SerialNo)", ""),
        ("Ordering number (OrderingNo)", ""),
        ("Firmware version (FirmwareVer)", ""),
        ("Product version (ProductVer)", ""),
        ("Production date (ProductionDate)", ""),
        ("Panel / cubicle location", ""),
        ("Main VT ratio (primary : secondary)", "_____ kV : _____ V"),
        ("Neutral VT ratio (primary : secondary)", "_____ V : _____ V"),
        ("Generator terminal CT ratio (primary : secondary)", "_____ A : _____ A"),
        ("Generator neutral CT ratio (primary : secondary)", "_____ A : _____ A"),
        ("Global base voltage (VBase)", "_____ kV"),
        ("Global base current (IBase)", "_____ A"),
        ("Active setting group", "GRP _____ (1–6)"),
        ("PCM600 project file name", ""),
        ("Notes", ""),
    ])

    doc.add_heading("1.4 Test Equipment", level=3)
    add_note(doc, "All injection equipment must be within calibration. Record "
             "calibration certificate numbers and due dates below.")
    make_table(doc,
               ["Equipment", "Make / Model", "Serial No.", "Last Cal. Date", "Cal. Due"],
               [["Three-phase test set", "Omicron CMC", "", "", ""],
                ["Laptop (PCM600)", "", "", "", ""],
                ["20 Hz injection set (64S/64R)", "ABB REX060 / equiv.", "", "", ""]],
               col_widths=[5.0, 4.0, 3.0, 2.7, 2.7], fixed_data_rows=2)

    doc.add_heading("1.5 Reference Drawings", level=3)
    make_table(doc, ["Drawing No.", "Title", "Revision", "Date"],
               [["", "Single line diagram", "", ""],
                ["", "AC schematic / CT-VT", "", ""],
                ["", "DC schematic / trip", "", ""],
                ["", "Binary I/O allocation", "", ""]],
               col_widths=[4.0, 7.0, 3.0, 3.0], fixed_data_rows=2)

    doc.add_heading("1.6 Test File & Reference Document Locations", level=3)
    add_para(doc, "Record the Omicron test file name and SharePoint locations for "
             "the key project documents (PCM600 project, settings, functional "
             "description, drawings).", italic=True, size=10)
    make_label_value_table(doc, [
        ("Omicron test file (.occ)", ""),
        ("PCM600 project file (.pcmp / .pcmi)", ""),
        ("SharePoint — Settings folder", ""),
        ("SharePoint — Functional description", ""),
        ("SharePoint — Drawings", ""),
    ])

    # ======================= SECTION 2 ===============================
    doc.add_heading("2. Safety, Isolations & Permits", level=1)
    add_para(doc, "Confirm the following before commencing any test work on the IED "
             "or its secondary circuits.", size=11)
    add_checks(doc, [
        "Permit to work / switching certificate received and read.",
        "Plant electrically isolated; circuit breaker(s) racked-out and locked-out (LOTO applied).",
        "VT secondary circuits isolated; VT MCBs OFF or links removed (mark each link removed).",
        "CT secondary circuits short-circuited at the test block before disconnection. Verify shorting.",
        "Trip and lockout outputs disabled / links removed to prevent unintended operation of plant.",
        "Trip lockout (86G) and master trip relays disabled where required. Record state below.",
        "64S / 64R sub-harmonic injection systems (REX060) blocked / isolated for testing if required.",
        "Adjacent live equipment identified; barriers / signs in place.",
        "Correct PPE: arc-rated clothing, insulating gloves, eye protection, safety footwear.",
        "Test leads inspected; no damaged insulation; correct ratings.",
        "Earth bond between test set and IED panel earth bar verified.",
        "Communications to SCADA / DCS (IEC 61850 / GOOSE) inhibited or operations notified to suppress alarms.",
    ])
    doc.add_heading("2.1 Isolation Record", level=3)
    make_table(doc, ["Isolation point", "Tag / Reference", "State left", "Time", "Initials"],
               [], col_widths=[5.0, 4.0, 3.0, 2.5, 2.5], fixed_data_rows=6)

    # ======================= SECTION 3 ===============================
    doc.add_heading("3. Pre-Test Checks", level=1)

    doc.add_heading("3.1 Visual Inspection", level=2)
    add_checks(doc, [
        "IED is mounted securely; no physical damage to fascia (LHMI) or terminals.",
        "LHMI display, push-buttons and LEDs all illuminate during lamp test.",
        "All wiring terminations tight (sample check); no signs of overheating or discolouration.",
        "No active alarm LEDs or messages on arrival. Record any present below.",
        "Auxiliary supply voltage at PSM terminals within tolerance (record value).",
        "INTERNAL FAIL contact wiring on the PSM intact and correctly terminated.",
    ])

    doc.add_heading("3.2 Connect to IED & Verify Status", level=2)
    add_checks(doc, [
        "Connect via LHMI (RJ-45 front port) or PCM600 via Ethernet (rear).",
        "Verify the “Normal” green status LED (above the LCD) is lit and steady — IED healthy.",
        "Check Main menu → Diagnostics → IED status → General: no FAIL or WARNING active.",
        "Record firmware version from Main menu → Diagnostics → IED status → Product identifiers.",
        "Record IEDProdType, ProductVer, SerialNo, OrderingNo, ProductionDate, FirmwareVer.",
        "Compare firmware version against the plant register.",
    ])
    make_label_value_table(doc, [
        ("“Normal” green LED state",
         "☐ Lit & steady (healthy)    ☐ Off / flashing (investigate)"),
        ("IED status — General", "☐ No FAIL / WARNING    ☐ Signal active (record below)"),
        ("Firmware version read", ""),
        ("Matches plant register?", "☐ Yes    ☐ No (raise finding)"),
    ])

    doc.add_heading("3.3 Self-Supervision Review  (ABB-specific)", level=2)
    add_note(doc, "The REG670 runs continuous hardware and software self-supervision. "
             "This is a key differentiator from the SEL-300G and MUST be reviewed every "
             "visit. The INTERRSIG function block reports the software status; the "
             "potential-free INTERNAL FAIL contact on the PSM reports the hardware status.")
    add_checks(doc, [
        "Navigate to Main menu → Diagnostics → Internal events (last 40 events, FIFO).",
        "Review and record ALL internal events since the last test visit.",
        "Confirm no active FAIL, WARNING, RTCERROR, TIMESYNCHERROR, RTEERROR, IEC61850ERROR, "
        "WATCHDOG, LMDERROR or APPERROR signals (Diagnostics → IED status → General).",
        "Check for module-specific errors: PSM-Error, ADOne-Error, BIM-Error, BOM-Error, "
        "IOM-Error, MIM-Error, LDCM-Error.",
        "Check the INTERNAL FAIL contact on the PSM is de-energised (i.e. IED healthy).",
        "Verify no alarm LEDs are lit that indicate a self-supervision issue.",
        "Record overall self-supervision status (All OK / active events noted).",
    ])
    make_table(doc, ["INTERRSIG signal", "State (Active / Inactive)", "Notes"],
               [["FAIL (Internal Fail)", "", ""],
                ["WARNING (Internal Warning)", "", ""],
                ["RTCERROR (RTC hardware)", "", ""],
                ["TIMESYNCHERROR (time sync lost)", "", ""],
                ["RTEERROR (runtime execution)", "", ""],
                ["IEC61850ERROR (61850 stack)", "", ""],
                ["WATCHDOG (heavy CPU >5 min)", "", ""],
                ["LMDERROR (LON interface)", "", ""],
                ["APPERROR (application thread)", "", ""]],
               col_widths=[6.5, 5.0, 5.5])
    make_label_value_table(doc, [
        ("INTERNAL FAIL contact (PSM)", "☐ De-energised (healthy)    ☐ Asserted (fault)"),
        ("Internal events since last visit", "☐ None    ☐ Recorded (attach list)"),
        ("Overall self-supervision status", "☐ All OK    ☐ Active events noted (raise finding)"),
    ])
    doc.add_heading("3.3.1 Internal Events Log (record since last visit)", level=3)
    make_table(doc, ["#", "Date / Time", "Internal event / signal", "Action taken"],
               [], col_widths=[1.5, 4.5, 6.5, 4.5], fixed_data_rows=8)

    doc.add_heading("3.4 Time / Clock Check", level=2)
    add_checks(doc, [
        "Navigate to Main menu → Diagnostics → IED status → General for RTC status.",
        "Compare IED clock with GPS / station reference time source.",
        "Note the TIMESYNCHERROR signal state.",
        "Record the configured time source (Main menu → Settings → Time).",
    ])
    add_note(doc, "The RTC is capacitor-backed. After ~5 days powered-off the clock may "
             "drift around 3 s/day and can reset. Record drift and correct at end of visit "
             "if outside tolerance (typically ±1 s).")
    make_label_value_table(doc, [
        ("Reference time source", "☐ GPS    ☐ IRIG-B    ☐ SNTP    ☐ LON    ☐ SPA    ☐ BIN"),
        ("IED time (read)", ""),
        ("Reference time", ""),
        ("Drift (s)", ""),
        ("TIMESYNCHERROR state", "☐ Inactive    ☐ Active"),
    ])

    doc.add_heading("3.5 Download Event Records", level=2)
    add_checks(doc, [
        "Export events via LHMI → Events or PCM600 Event Viewer.",
        "Save with naming convention: YYYYMMDD_[IED tag]_EVE_asfound.",
        "Review for unexpected pickups, operations or binary-input chatter since last visit.",
        "Cross-reference any abnormal events with operations / SCADA log.",
    ])

    doc.add_heading("3.6 Download Disturbance Records", level=2)
    add_checks(doc, [
        "LHMI → Disturbance records (DRPRDRE): download all records since last visit via PCM600.",
        "Save with naming convention: YYYYMMDD_[IED tag]_DIST_asfound.",
        "Review waveforms for any anomalies or unexplained operations.",
    ])

    doc.add_heading("3.7 Settings — As-Found", level=2)
    add_checks(doc, [
        "Open PCM600, connect to the IED and upload (read) the current settings.",
        "Save the PCM600 project as: asfound_[date]_[IED tag].pcm.",
        "Compare against the approved master settings.",
        "Document every difference. If unauthorised changes are found, STOP and notify the asset owner.",
    ])
    make_label_value_table(doc, [
        ("As-found project file name", ""),
        ("Master settings reference", ""),
        ("Differences found (Y/N)", ""),
    ])
    make_table(doc, ["#", "Setting / Parameter", "Master value", "As-found value", "Authorised? (Y/N)"],
               [], col_widths=[1.5, 6.0, 4.0, 4.0, 3.0], fixed_data_rows=6)

    doc.add_heading("3.8 Settings vs. Protection Functional Description", level=2)
    add_checks(doc, [
        "Open the latest approved Protection Functional Description for this asset.",
        "Verify enabled functions on the IED align with those required by the document.",
        "Spot-check key element settings (pickup, time delay, characteristic) against the document.",
        "Confirm trip matrix / SMBO output mapping matches the functional description.",
        "Confirm CT / VT ratios and VBase / IBase match nameplate / SLD and section 1.3 above.",
        "Record the active setting group (GRP 1–6).",
    ])
    make_label_value_table(doc, [
        ("Functional description rev.", ""),
        ("Active setting group", "GRP _____"),
        ("Discrepancies found (Y/N)", ""),
    ])

    # ======================= SECTION 4 ===============================
    doc.add_heading("4. Analogue Input Verification", level=1)
    add_para(doc, "The REG670 uses SMAI (Signal Matrix for Analog Inputs) preprocessing "
             "and reports PRIMARY values via the Measurements menu. Inject known secondary "
             "signals from the test set and read the PRIMARY value reported by the IED.", size=11)
    add_note(doc, "All measurements below are PRIMARY values read from Main menu → "
             "Measurements. Expected Primary = injected secondary × configured ratio "
             "(VBase / IBase scaling applies). The xlsx template computes Expected Primary, "
             "error % and pass/fail automatically.")

    doc.add_heading("4.1 Main VT Inputs — Three-Phase Voltage Injection", level=2)
    add_para(doc, "Inject balanced three-phase secondary voltage (L1/L2/L3). Read primary "
             "from Main menu → Measurements.", italic=True, size=10)
    add_checks(doc, [
        "Test block links / VT MCBs isolated; injection leads connected to IED terminals.",
        "Phase rotation set to ABC (or as installed).",
    ])
    ai_hdr = ["Step", "Injected (V sec)", "Angle (°)", "Phase / Input",
              "Expected Primary", "Relay Primary Reported", "Relay Angle (°)",
              "Mag Err %", "P/F"]
    ai_w = [1.3, 2.4, 1.8, 2.6, 2.6, 2.6, 2.0, 1.8, 1.5]
    make_table(doc, ai_hdr,
               [["1", "63.5", "0", "L1 (VA)", "", "", "", "", ""],
                ["1", "63.5", "-120", "L2 (VB)", "", "", "", "", ""],
                ["1", "63.5", "120", "L3 (VC)", "", "", "", "", ""],
                ["2", "30", "0", "L1 (VA)", "", "", "", "", ""],
                ["2", "30", "-120", "L2 (VB)", "", "", "", "", ""],
                ["2", "30", "120", "L3 (VC)", "", "", "", "", ""]],
               col_widths=ai_w)

    doc.add_heading("4.2 Neutral VT Input — Generator Neutral Voltage (VN)", level=2)
    add_para(doc, "Inject single-phase voltage into the dedicated neutral-voltage input.",
             italic=True, size=10)
    make_table(doc, ai_hdr,
               [["1", "10", "0", "VN", "", "", "", "", ""],
                ["2", "30", "0", "VN", "", "", "", "", ""],
                ["3", "60", "0", "VN", "", "", "", "", ""],
                ["4", "100", "0", "VN", "", "", "", "", ""]],
               col_widths=ai_w)

    doc.add_heading("4.3 Generator Terminal CT — Three-Phase Current Injection", level=2)
    add_para(doc, "Inject balanced three-phase current into the stator / terminal-side CT "
             "inputs (L1/L2/L3).", italic=True, size=10)
    add_checks(doc, [
        "CT shorting links applied at test block; relay-side connected to test set.",
        "Phase rotation ABC; angles per below.",
    ])
    make_table(doc, ai_hdr,
               [["1", "1", "0", "L1 terminal", "", "", "", "", ""],
                ["1", "1", "-120", "L2 terminal", "", "", "", "", ""],
                ["1", "1", "120", "L3 terminal", "", "", "", "", ""],
                ["2", "5", "0", "L1 terminal", "", "", "", "", ""],
                ["2", "5", "-120", "L2 terminal", "", "", "", "", ""],
                ["2", "5", "120", "L3 terminal", "", "", "", "", ""]],
               col_widths=ai_w)

    doc.add_heading("4.4 Generator Neutral CT — Three-Phase Current Injection", level=2)
    add_para(doc, "Inject balanced three-phase current into the neutral-side CT inputs "
             "(used by 87G GENPDIF).", italic=True, size=10)
    add_checks(doc, [
        "Confirm correct terminal block — neutral-side inputs.",
        "Polarity matches terminal-side injection so 87G operating current ≈ 0 with both sides equal.",
    ])
    make_table(doc, ai_hdr,
               [["1", "1", "0", "L1 neutral", "", "", "", "", ""],
                ["1", "1", "-120", "L2 neutral", "", "", "", "", ""],
                ["1", "1", "120", "L3 neutral", "", "", "", "", ""],
                ["2", "5", "0", "L1 neutral", "", "", "", "", ""],
                ["2", "5", "-120", "L2 neutral", "", "", "", "", ""],
                ["2", "5", "120", "L3 neutral", "", "", "", "", ""]],
               col_widths=ai_w)

    doc.add_heading("4.5 87G Differential Stability (Through-Current Injection)", level=2)
    add_para(doc, "Inject balanced three-phase through-current into both terminal-side and "
             "neutral-side CT inputs simultaneously. Verify differential current Id ≈ 0 and "
             "bias current Ibias ≈ rated. Read from Main menu → Measurements (GENPDIF).",
             italic=True, size=10)
    make_table(doc, ["Phase", "Terminal I (A pri)", "Neutral I (A pri)",
                     "Id operating (A)", "Ibias (A)", "Acceptable?"],
               [["L1", "", "", "", "", ""],
                ["L2", "", "", "", "", ""],
                ["L3", "", "", "", "", ""]],
               col_widths=[2.5, 3.2, 3.2, 3.0, 2.8, 2.8])

    # ======================= SECTION 5 ===============================
    doc.add_heading("5. Binary Input / Output Functional Tests", level=1)
    add_para(doc, "The REG670 maps physical I/O via SMBI (binary inputs) and SMBO (binary "
             "outputs) function blocks. Actual labels are site-specific — fill from the "
             "binary I/O allocation drawing.", size=11)

    doc.add_heading("5.1 Binary Inputs (SMBI)", level=2)
    add_para(doc, "Apply rated voltage to each input (or operate the source field device) "
             "and confirm assertion via LHMI / PCM600 signal monitoring.", italic=True, size=10)
    bi_defaults = [
        ("BI01", "86G lockout relay asserted"),
        ("BI02", "Plant available / off-line"),
        ("BI03", "Generator breaker open"),
        ("BI04", "Exciter breaker open"),
        ("BI05", "Test mode enable"),
        ("BI06", "REF695 block (if applied)"),
    ]
    make_table(doc, ["Input", "Description / Source (site drawing)", "Pickup V (V)",
                     "Asserted? (Y/N)", "Event captured? (Y/N)", "Initials"],
               [[bi, desc, "", "", "", ""] for bi, desc in bi_defaults],
               col_widths=[2.2, 6.0, 2.6, 2.8, 3.0, 2.0], fixed_data_rows=4)

    doc.add_heading("5.2 Binary Outputs (SMBO)", level=2)
    add_para(doc, "Force each output via PCM600 test / signal forcing, or by asserting the "
             "driving element. Verify continuity and target-device operation.", italic=True, size=10)
    add_note(doc, "Trip and lockout outputs MUST remain disabled at the panel / lockout "
             "relay if plant cannot tolerate an operation.")
    bo_defaults = [
        ("BO01", "Trip Generator Breaker (87G/87T/27/59)"),
        ("BO02", "Trip Exciter Breaker (40)"),
        ("BO03", "Trip & Lockout 86G"),
        ("BO04", "Alarm"),
        ("BO05", "INTERNAL FAIL (INTERRSIG FAIL)"),
    ]
    make_table(doc, ["Output", "Driven by / Function (site drawing)", "Force method",
                     "Continuity (Y/N)", "Operates target (Y/N)", "Initials"],
               [[bo, desc, "PCM600 force", "", "", ""] for bo, desc in bo_defaults],
               col_widths=[2.2, 6.0, 2.8, 2.8, 3.0, 2.0], fixed_data_rows=4)

    # ======================= SECTION 6 ===============================
    doc.add_heading("6. Protection Element Testing (Omicron)", level=1)

    doc.add_heading("6.1 Test File Structure", level=2)
    add_para(doc, "Inside the Omicron Test Universe (.occ) file for this asset, create one "
             "folder/section per protection element tested. Each should contain:", size=11)
    for it in ["A document block listing the as-applied settings (pickup, time delay, "
               "characteristic, restraint, blocking).",
               "Tester instructions: how to inject (currents, voltages, angles, fault "
               "types) and what to block to isolate the element.",
               "The test modules (ramp, shot, sequencer, etc.) configured for that element.",
               "A results table recording measured pickup, drop-out, operate time and pass/fail."]:
        doc.add_paragraph(it, style="List Bullet")
    add_note(doc, "Keep folder naming consistent — recommended “<ANSI>_<ABB function>” "
             "(e.g. 87G_GENPDIF, 40_LEXPDIS).")

    doc.add_heading("6.2 Element Selection Strategy", level=2)
    add_para(doc, "It is not necessary to test every enabled element on each visit. Rotate "
             "elements so coverage is achieved over time. Prioritise:", size=11)
    for it in ["Elements not tested in the longest time.",
               "Elements re-set or modified since last test.",
               "Elements that operated in service since last test.",
               "At least one critical element per visit (e.g. 87G, 40).",
               "Avoid testing only the easy elements; alternate to give all elements a turn."]:
        doc.add_paragraph(it, style="List Bullet")

    doc.add_heading("6.3 Element Tracker — Last Tested / This Visit Selection", level=2)
    elements = [
        ("87G", "GENPDIF", "GENPDIF", "Generator differential"),
        ("87T", "T2WPDIF", "T2WPDIF", "Unit transformer differential"),
        ("87N", "REFPDIF", "REFPDIF", "Restricted earth fault (low-Z)"),
        ("64S", "STTIPHIZ", "STTIPHIZ", "100% stator earth fault (20 Hz inj.)"),
        ("64R", "ROTIPHIZ", "ROTIPHIZ", "Rotor earth fault (injection)"),
        ("40", "LEXPDIS", "LEXPDIS", "Loss of excitation (impedance)"),
        ("78", "PSPPPAM", "PSPPPAM", "Pole slip (lens characteristic)"),
        ("51V", "VRPVOC", "VRPVOC", "Volt-restrained overcurrent (stator backup)"),
        ("50/51", "OC4PTOC", "OC4PTOC", "Phase overcurrent backup"),
        ("50N/51N", "EF4PTOC", "EF4PTOC", "Residual overcurrent backup"),
        ("46", "NS4PTOC", "NS4PTOC", "Negative-sequence overcurrent"),
        ("49S", "GSPTTR", "GSPTTR", "Stator thermal overload (I²t)"),
        ("49R", "GRPTTR", "GRPTTR", "Rotor thermal overload (I²t)"),
        ("27", "UV2PTUV", "UV2PTUV", "Undervoltage (2-step)"),
        ("59", "OV2PTOV", "OV2PTOV", "Overvoltage (2-step)"),
        ("59N", "ROV2PTOV", "ROV2PTOV", "Residual / zero-seq. overvoltage"),
        ("24", "OEXPVPH", "OEXPVPH", "Overexcitation (V/Hz)"),
        ("32", "GOPPDOP", "GOPPDOP", "Directional overpower / reverse power"),
        ("37", "GUPPDUP", "GUPPDUP", "Directional underpower"),
        ("50BF", "CCRBRF", "CCRBRF", "Breaker failure (3-phase)"),
        ("52PD", "CCPDSC", "CCPDSC", "CB pole discrepancy"),
        ("25", "SESRSYN", "SESRSYN", "Synchronism / energising check"),
        ("50AE", "AEGPVOC", "AEGPVOC", "Accidental / inadvertent energising"),
        ("81U", "SAPTUF", "SAPTUF", "Underfrequency"),
        ("81O", "SAPTOF", "SAPTOF", "Overfrequency"),
        ("81R", "SAPFRC", "SAPFRC", "Rate of change of frequency (ROCOF)"),
        ("60", "FUFSPVC", "FUFSPVC", "Fuse / VT failure supervision"),
        ("CCS", "CCS", "CCS", "Current circuit supervision"),
    ]
    rows = [[ansi, abb, desc, "", "", "", ""] for ansi, abb, _id, desc in elements]
    make_table(doc, ["ANSI", "ABB Function", "Description", "Enabled? (Y/N)",
                     "Last tested (date)", "Test this visit? (Y/N)", "Result"],
               rows, col_widths=[1.8, 2.6, 6.4, 1.8, 2.4, 2.2, 1.8])

    doc.add_heading("6.4 Per-Element Test Result Template", level=2)
    add_para(doc, "Copy this template into the test record for each element tested in this visit.",
             italic=True, size=10)
    make_label_value_table(doc, [
        ("Element / ANSI", ""),
        ("ABB function (IEC 61850 ID)", ""),
        ("Setting tested (pickup / reach / curve)", ""),
        ("Pickup setting", ""),
        ("Pickup measured", ""),
        ("Pickup error %", ""),
        ("Time delay setting (s)", ""),
        ("Time measured (s)", ""),
        ("Time error %", ""),
        ("Result", "☐ PASS    ☐ FAIL"),
    ])

    # ======================= SECTION 7 ===============================
    doc.add_heading("7. Post-Test Checks", level=1)

    doc.add_heading("7.1 Re-compare Settings to Master", level=2)
    add_checks(doc, [
        "Read settings from the IED again after element testing complete (PCM600 upload).",
        "Save the PCM600 project as: asleft_[date]_[IED tag].pcm.",
        "Compare as-left against the approved master settings.",
        "Confirm zero unintended differences (only authorised changes, if any, remain).",
        "Confirm any temporarily disabled / modified functions have been restored.",
    ])
    make_label_value_table(doc, [
        ("As-left project file name", ""),
        ("Compare result vs. master", "☐ Identical    ☐ Authorised changes only    ☐ Differences (raise finding)"),
    ])

    doc.add_heading("7.2 Clear Events / Reset Targets", level=2)
    add_checks(doc, [
        "Clear test-mode alarms via LHMI → Clear.",
        "Reset LHMI alarm LEDs and any latched indications.",
        "Confirm Internal events / disturbance buffers reviewed and archived as required.",
    ])

    doc.add_heading("7.3 Time Synchronisation Check", level=2)
    add_checks(doc, [
        "Re-read IED date/time (Diagnostics → IED status → General).",
        "If drift was outside tolerance, correct via the configured time source.",
        "Confirm time synchronisation healthy: TIMESYNCHERROR inactive.",
        "Re-read time and confirm within tolerance of reference.",
    ])
    make_label_value_table(doc, [
        ("Time source", "☐ GPS    ☐ IRIG-B    ☐ SNTP    ☐ LON    ☐ SPA    ☐ BIN"),
        ("Final IED time vs reference", ""),
    ])

    doc.add_heading("7.4 Restore Plant", level=2)
    add_checks(doc, [
        "Disconnect all test leads from IED terminals.",
        "Replace all test block links / reinstate VT MCBs.",
        "Remove CT shorting links and confirm CT secondary continuity through to source.",
        "Restore trip and lockout output links / re-enable trips.",
        "Confirm the self-supervision “Normal” (green) LED is re-lit and steady after test completion.",
        "Clear any remaining test-mode alarms via LHMI → Clear.",
        "Verify the INTERNAL FAIL contact on the PSM is de-energised (healthy).",
        "Confirm the active setting group is restored to the correct group (NOT left in a test group).",
        "Confirm 64S / 64R injection systems (REX060 units) are re-enabled if blocked for testing.",
        "Verify no targets / alarm LEDs remain active on the IED or panel.",
        "Confirm SCADA / DCS analogue values match IED metering (live readings, post-restoration).",
        "Re-energise VTs; observe sensible secondary voltages on the IED before declaring complete.",
        "Operations / control room notified that work is complete.",
        "Permit returned and isolations removed (LOTO removed).",
    ])

    doc.add_heading("7.5 Final Self-Supervision Check", level=2)
    add_checks(doc, [
        "Confirm Main menu → Diagnostics → IED status → General reports no FAIL or WARNING.",
        "Confirm the “Normal” green LED is lit and steady.",
        "Confirm the IED is in service / not in test mode per plant requirements.",
        "Capture final screenshots of IED status / Measurements / LED pages for the record.",
    ])

    # ======================= SECTION 8 ===============================
    doc.add_heading("8. Findings, Defects & Recommendations", level=1)
    add_para(doc, "Use this log to record every individual finding raised during the test. "
             "The cover-page summary should reference these items.", size=11)
    make_table(doc, ["#", "Finding / Defect", "Severity", "Action / NCR raised", "Closed?"],
               [], col_widths=[1.5, 8.0, 3.0, 4.0, 2.5], fixed_data_rows=8)

    # ======================= SECTION 9 ===============================
    doc.add_heading("9. Test Files Index", level=1)
    add_para(doc, "Record full file paths / network locations for all data captured during the test.",
             size=11)
    make_table(doc, ["File type", "File name", "Location", "Initials"],
               [["Events (EVE)", "", "", ""],
                ["Disturbance records (DRPRDRE)", "", "", ""],
                ["Internal events list", "", "", ""],
                ["As-found PCM600 project", "", "", ""],
                ["As-left PCM600 project", "", "", ""],
                ["Omicron test file (.occ)", "", "", ""],
                ["Screenshots", "", "", ""]],
               col_widths=[5.0, 5.0, 6.0, 2.5], fixed_data_rows=3)

    # ======================= SECTION 10 ==============================
    doc.add_heading("10. Test Sign-Off", level=1)
    add_para(doc, "By signing below, the tester and witness confirm that all activities "
             "recorded above were performed and the results are accurate.", size=11)
    make_label_value_table(doc, [
        ("Tester name", ""),
        ("Tester signature", ""),
        ("Date", ""),
        ("Witness name", ""),
        ("Witness signature", ""),
        ("Date", ""),
        ("Authorising engineer (acceptance)", ""),
        ("Signature", ""),
        ("Date", ""),
        ("Overall test result", "☐ PASS    ☐ PASS with defects    ☐ FAIL"),
    ])

    out = os.path.join(HERE, "REG670_Test_Check_Sheet.docx")
    doc.save(out)
    return out


# ===========================================================================
#  EXCEL RESULTS TEMPLATE
# ===========================================================================

def build_xlsx():
    wb = openpyxl.Workbook()

    # styles
    f_title = Font(name="Calibri", size=18, bold=True, color=BRAND)
    f_h = Font(name="Calibri", size=12, bold=True, color=WHITE)
    f_colhdr = Font(name="Calibri", size=10, bold=True, color=WHITE)
    f_label = Font(name="Calibri", size=10, bold=True, color=BLACK)
    f_body = Font(name="Calibri", size=10, color=BLACK)
    f_input = Font(name="Calibri", size=10, color=INPUT_BLUE)
    f_note = Font(name="Calibri", size=9, italic=True, color="595959")

    fill_brand = PatternFill("solid", fgColor=BRAND)
    fill_accent = PatternFill("solid", fgColor=ACCENT)
    fill_note = PatternFill("solid", fgColor=NOTE_BG)
    fill_white = PatternFill("solid", fgColor=WHITE)

    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    ctr = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ctrw = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    def hdr_cell(ws, coord, text, fill=fill_brand, font=f_colhdr, align=ctr):
        c = ws[coord]; c.value = text; c.fill = fill; c.font = font
        c.alignment = align; c.border = border
        return c

    def section_title(ws, coord, text):
        c = ws[coord]; c.value = text; c.font = f_h; c.fill = fill_brand
        c.alignment = left
        return c

    def label(ws, coord, text):
        c = ws[coord]; c.value = text; c.font = f_label; c.fill = fill_accent
        c.alignment = left; c.border = border
        return c

    def inputcell(ws, coord, val=None):
        c = ws[coord]; c.value = val; c.font = f_input; c.fill = fill_white
        c.alignment = left; c.border = border
        return c

    def col_widths(ws, widths):
        for col, w in widths.items():
            ws.column_dimensions[col].width = w

    PASS_RULE = lambda: CellIsRule(operator="equal", formula=['"PASS"'],
                                   font=Font(color=PASS_TXT, bold=True),
                                   fill=PatternFill("solid", fgColor=PASS_FILL))
    FAIL_RULE = lambda: CellIsRule(operator="equal", formula=['"FAIL"'],
                                   font=Font(color=FAIL_TXT, bold=True),
                                   fill=PatternFill("solid", fgColor=FAIL_FILL))

    # ----------------------------------------------------------------
    # SHEET 1: Cover
    # ----------------------------------------------------------------
    cov = wb.active
    cov.title = "1. Cover"
    cov.sheet_view.showGridLines = False
    col_widths(cov, {"A": 2, "B": 34, "C": 22, "D": 16, "E": 14, "F": 14, "G": 10})

    cov["D2"] = "PROTECTION RELAY TEST RESULTS"; cov["D2"].font = f_title
    cov["D3"] = "ABB Relion REG670 2.0 (ANSI) Generator Protection IED — Live Results Template"
    cov["D3"].font = Font(name="Calibri", size=11, italic=True, color=BRAND)

    section_title(cov, "B4", "Test Summary & Findings (filled in at end of visit)")
    summary = [("B5", "Overall outcome"), ("B6", "Test completion date"),
               ("B7", "Summary & Findings"),
               ("B8", "Self-supervision status on departure")]
    for coord, txt in summary:
        label(cov, coord, txt)
        inputcell(cov, "C" + coord[1:])

    section_title(cov, "B14", "Test Information")
    info = [("Site / Power station", ""),
            ("Asset identifier (generator / unit)", ""),
            ("Work order / Job number", ""),
            ("Reason for test", "Routine"),
            ("Test date (start)", ""),
            ("Test date (finish)", "")]
    for i, (lab, val) in enumerate(info):
        r = 15 + i
        label(cov, f"B{r}", lab)
        inputcell(cov, f"C{r}", val if val else None)

    section_title(cov, "B22", "Personnel")
    pers = ["Tester (name)", "Tester signature", "Witness / 2nd person",
            "Witness signature", "Authorising engineer"]
    for i, lab in enumerate(pers):
        r = 23 + i
        label(cov, f"B{r}", lab); inputcell(cov, f"C{r}")

    section_title(cov, "B29", "IED Under Test")
    iut = [("IED model", "ABB Relion REG670 2.0 (ANSI)"),
           ("IED tag / functional ID", ""),
           ("Serial number (SerialNo)", ""),
           ("Ordering number (OrderingNo)", ""),
           ("Firmware version (FirmwareVer)", ""),
           ("Product version (ProductVer)", ""),
           ("Production date (ProductionDate)", ""),
           ("Panel / cubicle location", ""),
           ("Active setting group (GRP 1–6)", ""),
           ("Global base voltage VBase (kV)", ""),
           ("Global base current IBase (A)", ""),
           ("PCM600 project file name", ""),
           ("Notes", "")]
    for i, (lab, val) in enumerate(iut):
        r = 30 + i
        label(cov, f"B{r}", lab); inputcell(cov, f"C{r}", val if val else None)

    # ---- Ratios block (drives Expected Primary) ----
    rrow = 44
    section_title(cov, f"B{rrow}",
                  "CT / VT Ratios (primary : secondary)  —  drives Expected Primary in analogue tables")
    rrow += 1
    for j, txt in zip("BCDEF", ["Ratio", "Primary", "Secondary", "Ratio (P/S)", "Unit"]):
        hdr_cell(cov, f"{j}{rrow}", txt)
    ratio_rows = [
        ("Main VT (line-N)", 11000, 63.5, "V"),
        ("Neutral VT (VN)", 11000, 110, "V"),
        ("Generator terminal CT", 1000, 5, "A"),
        ("Generator neutral CT", 1000, 5, "A"),
    ]
    ratio_first = rrow + 1
    for i, (lab, pri, sec, unit) in enumerate(ratio_rows):
        r = ratio_first + i
        label(cov, f"B{r}", lab)
        inputcell(cov, f"C{r}", pri)
        inputcell(cov, f"D{r}", sec)
        ec = cov[f"E{r}"]; ec.value = f'=IFERROR(C{r}/D{r},"")'
        ec.font = f_body; ec.alignment = ctr; ec.border = border
        ec.fill = PatternFill("solid", fgColor="F2F2F2")
        uc = cov[f"F{r}"]; uc.value = unit; uc.font = f_body; uc.alignment = ctr; uc.border = border
    # named ranges for ratios
    rmap = {
        "Ratio_VTmain": f"E{ratio_first}", "Ratio_VTmain_pri": f"C{ratio_first}",
        "Ratio_VTmain_sec": f"D{ratio_first}",
        "Ratio_VTneut": f"E{ratio_first+1}", "Ratio_VTneut_pri": f"C{ratio_first+1}",
        "Ratio_VTneut_sec": f"D{ratio_first+1}",
        "Ratio_CTterm": f"E{ratio_first+2}", "Ratio_CTterm_pri": f"C{ratio_first+2}",
        "Ratio_CTterm_sec": f"D{ratio_first+2}",
        "Ratio_CTneut": f"E{ratio_first+3}", "Ratio_CTneut_pri": f"C{ratio_first+3}",
        "Ratio_CTneut_sec": f"D{ratio_first+3}",
    }

    # ---- Tolerances ----
    trow = ratio_first + 4 + 1
    section_title(cov, f"B{trow}", "Test Tolerances (drives auto pass/fail)")
    tol_rows = [
        ("Voltage magnitude tolerance (%)", 1, "TolV_pct"),
        ("Current magnitude tolerance (%)", 1, "TolI_pct"),
        ("Phase angle tolerance (°)", 1, "TolAngle_deg"),
        ("Element pickup tolerance (%)", 5, "TolPickup_pct"),
        ("Element timing tolerance (%)", 5, "TolTime_pct"),
    ]
    tol_first = trow + 1
    tolmap = {}
    for i, (lab, val, name) in enumerate(tol_rows):
        r = tol_first + i
        label(cov, f"B{r}", lab)
        c = cov[f"C{r}"]; c.value = val; c.font = f_input
        c.fill = fill_note; c.alignment = ctr; c.border = border
        tolmap[name] = f"C{r}"

    # ---- roll-up ----
    urow = tol_first + 5 + 1
    section_title(cov, f"B{urow}", "Auto-rolled-up status (calculated)")
    ai_cell = f"C{urow+1}"; el_cell = f"C{urow+2}"; ov_cell = f"C{urow+3}"
    label(cov, f"B{urow+1}", "Analogue inputs")
    c = cov[ai_cell]
    c.value = ('=IF(COUNTIF(\'4. Analogue Inputs\'!$K$1:$K$300,"FAIL")>0,"FAIL",'
               'IF(COUNTIF(\'4. Analogue Inputs\'!$K$1:$K$300,"PASS")>0,"PASS",""))')
    c.font = f_label; c.alignment = ctr; c.border = border
    label(cov, f"B{urow+2}", "Element tests")
    c = cov[el_cell]
    c.value = ('=IF(COUNTIF(\'7. Element Results\'!$L$1:$L$300,"FAIL")>0,"FAIL",'
               'IF(COUNTIF(\'7. Element Results\'!$L$1:$L$300,"PASS")>0,"PASS",""))')
    c.font = f_label; c.alignment = ctr; c.border = border
    label(cov, f"B{urow+3}", "Auto overall")
    c = cov[ov_cell]
    c.value = (f'=IF(OR({ai_cell}="FAIL",{el_cell}="FAIL"),"FAIL",'
               f'IF(AND({ai_cell}="PASS",{el_cell}="PASS"),"PASS","INCOMPLETE"))')
    c.font = f_label; c.alignment = ctr; c.border = border

    # conditional formatting on roll-up cells
    for cc in (ai_cell, el_cell, ov_cell):
        cov.conditional_formatting.add(cc, PASS_RULE())
        cov.conditional_formatting.add(cc, FAIL_RULE())

    lrow = urow + 5
    cov[f"B{lrow}"] = ("Legend: blue text = user input.  Grey/auto cells are calculated.  "
                       "Yellow tolerance cells drive PASS/FAIL.")
    cov[f"B{lrow}"].font = f_note

    # remember overall cell for sign-off reference
    overall_ref = f"'1. Cover'!C5"

    # ----------------------------------------------------------------
    # SHEET 2: Equip & Refs
    # ----------------------------------------------------------------
    eq = wb.create_sheet("2. Equip & Refs")
    eq.sheet_view.showGridLines = False
    col_widths(eq, {"A": 2, "B": 26, "C": 20, "D": 16, "E": 16, "F": 14})
    eq["B2"] = "Test Equipment & Reference Drawings"; eq["B2"].font = f_title
    section_title(eq, "B4", "Test Equipment")
    for j, txt in zip("BCDEF", ["Equipment", "Make / Model", "Serial No.", "Last Cal.", "Cal. Due"]):
        hdr_cell(eq, f"{j}4", txt)
    eqrows = [("Three-phase test set", "Omicron CMC"),
              ("Laptop (PCM600)", ""),
              ("20 Hz injection set (64S/64R)", "ABB REX060"),
              ("", ""), ("", "")]
    for i, (a, b) in enumerate(eqrows):
        r = 5 + i
        for j, val in zip("BCDEF", [a, b, "", "", ""]):
            inputcell(eq, f"{j}{r}", val if val else None)

    section_title(eq, "B12", "Reference Drawings")
    for j, txt in zip("BCDE", ["Drawing No.", "Title", "Revision", "Date"]):
        hdr_cell(eq, f"{j}12", txt)
    drw = ["Single line diagram", "AC schematic / CT-VT", "DC schematic / trip",
           "Binary I/O allocation", "", ""]
    for i, t in enumerate(drw):
        r = 13 + i
        inputcell(eq, f"B{r}"); inputcell(eq, f"C{r}", t if t else None)
        inputcell(eq, f"D{r}"); inputcell(eq, f"E{r}")

    section_title(eq, "B21", "Test File & Reference Document Locations")
    locs = ["Omicron test file (.occ)", "PCM600 project file (.pcmp/.pcmi)",
            "SharePoint — Settings", "SharePoint — Functional description",
            "SharePoint — Drawings"]
    for i, lab in enumerate(locs):
        r = 22 + i
        label(eq, f"B{r}", lab); inputcell(eq, f"C{r}")

    # ----------------------------------------------------------------
    # SHEET 3: Pre-Test
    # ----------------------------------------------------------------
    pt = wb.create_sheet("3. Pre-Test")
    pt.sheet_view.showGridLines = False
    col_widths(pt, {"A": 2, "B": 6, "C": 60, "D": 12, "E": 26, "F": 10})
    pt["B2"] = "Pre-Test Checklist"; pt["B2"].font = f_title

    def checklist_block(ws, start_row, title, items, num_start):
        section_title(ws, f"B{start_row}", title)
        hr = start_row + 1
        for j, txt in zip("BCDEF", ["#", "Check", "Status", "Notes / Value", "Initials"]):
            hdr_cell(ws, f"{j}{hr}", txt)
        for i, it in enumerate(items):
            r = hr + 1 + i
            c = ws[f"B{r}"]; c.value = num_start + i; c.font = f_body; c.alignment = ctr; c.border = border
            c2 = ws[f"C{r}"]; c2.value = it; c2.font = f_body; c2.alignment = left; c2.border = border
            inputcell(ws, f"D{r}"); inputcell(ws, f"E{r}"); inputcell(ws, f"F{r}")
        return hr + 1 + len(items)

    n = 1
    r = checklist_block(pt, 4, "3.1 Visual Inspection", [
        "IED mounted securely; no physical damage to LHMI or terminals.",
        "LHMI display, buttons and LEDs illuminate during lamp test.",
        "All wiring terminations tight (sample); no overheating.",
        "No active alarm LEDs or messages on arrival.",
        "Aux supply at PSM within tolerance (record value).",
        "INTERNAL FAIL contact wiring on PSM intact.",
    ], n); n += 6
    r = checklist_block(pt, r + 1, "3.2 Connect to IED & Verify Status", [
        "Connected via LHMI (front RJ-45) or PCM600 (Ethernet).",
        '"Normal" green status LED lit and steady (IED healthy).',
        "Diagnostics > IED status > General: no FAIL/WARNING active.",
        "Firmware version recorded from Product identifiers (Cover).",
        "Firmware matches plant register.",
    ], n); n += 5
    r = checklist_block(pt, r + 1, "3.3 Self-Supervision Review (ABB-specific)", [
        "Diagnostics > Internal events reviewed; events since last visit recorded.",
        "No active FAIL / WARNING / RTCERROR / TIMESYNCHERROR signals.",
        "No RTEERROR / IEC61850ERROR / WATCHDOG / LMDERROR / APPERROR.",
        "No module errors (PSM/ADOne/BIM/BOM/IOM/MIM/LDCM).",
        "INTERNAL FAIL contact on PSM de-energised (healthy).",
        "No self-supervision alarm LEDs lit.",
        "Overall self-supervision status recorded (All OK / events noted).",
    ], n); n += 7
    r = checklist_block(pt, r + 1, "3.4 Time / Clock Check", [
        "Read IED date/time (Diagnostics > IED status > General); record drift.",
        "Time source identified (GPS / IRIG-B / SNTP / LON / SPA / BIN).",
        "TIMESYNCHERROR signal state noted.",
    ], n); n += 3
    r = checklist_block(pt, r + 1, "3.5 Event & Disturbance Records", [
        "Events exported (LHMI > Events / PCM600) as YYYYMMDD_[tag]_EVE_asfound.",
        "Disturbance records (DRPRDRE) downloaded as YYYYMMDD_[tag]_DIST_asfound.",
        "Reviewed for unexpected operations / pickups / chatter.",
        "Anomalies cross-referenced with operations / SCADA log.",
    ], n); n += 4
    r = checklist_block(pt, r + 1, "3.6 Settings — As-Found", [
        "PCM600 upload; saved as asfound_[date]_[tag].pcm.",
        "Compared as-found vs. master settings.",
        "All differences documented.",
        "Differences are authorised changes only.",
    ], n); n += 4
    r = checklist_block(pt, r + 1, "3.7 Settings vs. Functional Description", [
        "Enabled functions match functional description.",
        "Key element settings match document.",
        "CT/VT ratios and VBase/IBase match nameplate/SLD and Cover.",
        "Trip matrix / SMBO output mapping matches functional description.",
        "Active setting group recorded.",
    ], n); n += 5

    # ----------------------------------------------------------------
    # SHEET 4: Analogue Inputs  (the formula engine)
    # ----------------------------------------------------------------
    ai = wb.create_sheet("4. Analogue Inputs")
    ai.sheet_view.showGridLines = False
    col_widths(ai, {"A": 2, "B": 6, "C": 13, "D": 12, "E": 16, "F": 18,
                    "G": 16, "H": 12, "I": 12, "J": 12, "K": 8, "L": 24})
    ai["B2"] = "Analogue Input Verification — PRIMARY values"; ai["B2"].font = f_title
    ai["B3"] = ("Inject secondary; the IED reports PRIMARY values via Measurements using the "
                "ratios on the Cover. Expected Primary, error % and pass/fail compute automatically. "
                "Magnitude tolerance: TolV_pct / TolI_pct; angle tolerance: TolAngle_deg.")
    ai["B3"].font = f_note; ai["B3"].alignment = left
    ai["B4"] = "If Expected Primary is blank, check the corresponding CT/VT ratio is filled on the Cover."
    ai["B4"].font = f_note; ai["B4"].alignment = left

    ai_cols = ["Step", "Injected (sec)", "Inj Angle (°)", "Phase / Input",
               "Expected Primary", "Relay Primary Reported", "Relay Angle (°)",
               "Mag Err (%)", "Angle Err (°)", "P / F", "Notes"]

    def ai_block(start_row, title, ratio_name, tol_name, data):
        """data = list of (step, injected, angle, phase). Returns next free row + pf range."""
        section_title(ai, f"B{start_row}", title)
        hr = start_row + 1
        for j, txt in zip("BCDEFGHIJKL", ai_cols):
            hdr_cell(ai, f"{j}{hr}", txt)
        first = hr + 1
        for i, (step, inj, ang, phase) in enumerate(data):
            r = first + i
            inputcell(ai, f"B{r}", step); ai[f"B{r}"].alignment = ctr
            inputcell(ai, f"C{r}", inj); ai[f"C{r}"].alignment = ctr
            inputcell(ai, f"D{r}", ang); ai[f"D{r}"].alignment = ctr
            inputcell(ai, f"E{r}", phase)
            # Expected primary (calculated)
            fc = ai[f"F{r}"]; fc.value = f'=IFERROR(C{r}*{ratio_name},"")'
            fc.font = f_body; fc.alignment = ctr; fc.border = border
            fc.fill = PatternFill("solid", fgColor="F2F2F2")
            inputcell(ai, f"G{r}"); ai[f"G{r}"].alignment = ctr   # relay primary reported
            inputcell(ai, f"H{r}"); ai[f"H{r}"].alignment = ctr   # relay angle
            # mag err %
            ic = ai[f"I{r}"]
            ic.value = f'=IFERROR(IF(OR(F{r}="",G{r}=""),"",(G{r}-F{r})/F{r}*100),"")'
            ic.font = f_body; ic.alignment = ctr; ic.border = border
            ic.fill = PatternFill("solid", fgColor="F2F2F2")
            # angle err
            jc = ai[f"J{r}"]
            jc.value = f'=IFERROR(IF(OR(D{r}="",H{r}=""),"",H{r}-D{r}),"")'
            jc.font = f_body; jc.alignment = ctr; jc.border = border
            jc.fill = PatternFill("solid", fgColor="F2F2F2")
            # P/F. N() coerces blank "" to 0 so ABS never sees text (AND does not
            # short-circuit in all engines); ISNUMBER drives the blank-row test.
            kc = ai[f"K{r}"]
            kc.value = (f'=IF(AND(NOT(ISNUMBER(I{r})),NOT(ISNUMBER(J{r}))),"",'
                        f'IF(OR(AND(ISNUMBER(I{r}),ABS(N(I{r}))>{tol_name}),'
                        f'AND(ISNUMBER(J{r}),ABS(N(J{r}))>TolAngle_deg)),"FAIL","PASS"))')
            kc.font = f_body; kc.alignment = ctr; kc.border = border
            inputcell(ai, f"L{r}")
        last = first + len(data) - 1
        ai.conditional_formatting.add(f"K{first}:K{last}", PASS_RULE())
        ai.conditional_formatting.add(f"K{first}:K{last}", FAIL_RULE())
        return last + 2

    r = ai_block(6, "4.1 Main VT — Three-Phase Voltage Injection", "Ratio_VTmain", "TolV_pct",
                 [(1, 63.5, 0, "L1 (VA)"), (1, 63.5, -120, "L2 (VB)"), (1, 63.5, 120, "L3 (VC)"),
                  (2, 30, 0, "L1 (VA)"), (2, 30, -120, "L2 (VB)"), (2, 30, 120, "L3 (VC)")])
    r = ai_block(r, "4.2 Neutral VT — Generator Neutral Voltage (VN)", "Ratio_VTneut", "TolV_pct",
                 [(1, 10, 0, "VN"), (2, 30, 0, "VN"), (3, 60, 0, "VN"), (4, 100, 0, "VN")])
    r = ai_block(r, "4.3 Generator Terminal CT — Three-Phase Current Injection", "Ratio_CTterm", "TolI_pct",
                 [(1, 1, 0, "L1 terminal"), (1, 1, -120, "L2 terminal"), (1, 1, 120, "L3 terminal"),
                  (2, 5, 0, "L1 terminal"), (2, 5, -120, "L2 terminal"), (2, 5, 120, "L3 terminal")])
    r = ai_block(r, "4.4 Generator Neutral CT — Three-Phase Current Injection", "Ratio_CTneut", "TolI_pct",
                 [(1, 1, 0, "L1 neutral"), (1, 1, -120, "L2 neutral"), (1, 1, 120, "L3 neutral"),
                  (2, 5, 0, "L1 neutral"), (2, 5, -120, "L2 neutral"), (2, 5, 120, "L3 neutral")])
    # 4.5 stability (manual, no auto P/F)
    section_title(ai, f"B{r}", "4.5 87G Differential Stability (Through-Current Injection)")
    hr = r + 1
    for j, txt in zip("BCDEF", ["Phase", "Terminal I (A pri)", "Neutral I (A pri)",
                                "Id operating (A)", "Ibias (A)"]):
        hdr_cell(ai, f"{j}{hr}", txt)
    for i, ph in enumerate(["L1", "L2", "L3"]):
        rr = hr + 1 + i
        c = ai[f"B{rr}"]; c.value = ph; c.font = f_body; c.alignment = ctr; c.border = border
        for j in "CDEF":
            inputcell(ai, f"{j}{rr}"); ai[f"{j}{rr}"].alignment = ctr
    ai.freeze_panes = "B6"

    # ----------------------------------------------------------------
    # SHEET 5: Binary IO
    # ----------------------------------------------------------------
    bio = wb.create_sheet("5. Binary IO")
    bio.sheet_view.showGridLines = False
    col_widths(bio, {"A": 2, "B": 10, "C": 36, "D": 14, "E": 14, "F": 16, "G": 10, "H": 24, "I": 10})
    bio["B2"] = "Binary Input / Output Functional Tests (SMBI / SMBO)"; bio["B2"].font = f_title
    section_title(bio, "B4", "5.1 Binary Inputs (SMBI)")
    for j, txt in zip("BCDEFGHI", ["Input", "Description / Source", "Pickup V (V)",
                                   "Asserted? (Y/N)", "Event captured? (Y/N)", "Status",
                                   "Notes", "Initials"]):
        hdr_cell(bio, f"{j}5", txt)
    bin_in = [("BI01", "86G lockout relay asserted"), ("BI02", "Plant available / off-line"),
              ("BI03", "Generator breaker open"), ("BI04", "Exciter breaker open"),
              ("BI05", "Test mode enable"), ("BI06", "REF695 block (if applied)"),
              ("BI07", ""), ("BI08", "")]
    for i, (a, b) in enumerate(bin_in):
        r = 6 + i
        c = bio[f"B{r}"]; c.value = a; c.font = f_label; c.fill = fill_accent; c.alignment = ctr; c.border = border
        for j, v in zip("CDEFGHI", [b, None, None, None, None, None, None]):
            inputcell(bio, f"{j}{r}", v if v else None)

    section_title(bio, "B16", "5.2 Binary Outputs (SMBO)")
    for j, txt in zip("BCDEFGHI", ["Output", "Driven by / Function", "Force method",
                                   "Continuity (Y/N)", "Operates target (Y/N)", "Status",
                                   "Notes", "Initials"]):
        hdr_cell(bio, f"{j}17", txt)
    bin_out = [("BO01", "Trip Generator Breaker"), ("BO02", "Trip Exciter Breaker (40)"),
               ("BO03", "Trip & Lockout 86G"), ("BO04", "Alarm"),
               ("BO05", "INTERNAL FAIL (INTERRSIG FAIL)"), ("BO06", ""), ("BO07", ""), ("BO08", "")]
    for i, (a, b) in enumerate(bin_out):
        r = 18 + i
        c = bio[f"B{r}"]; c.value = a; c.font = f_label; c.fill = fill_accent; c.alignment = ctr; c.border = border
        inputcell(bio, f"C{r}", b if b else None)
        inputcell(bio, f"D{r}", "PCM600 force")
        for j in "EFGHI":
            inputcell(bio, f"{j}{r}")

    # ----------------------------------------------------------------
    # SHEET 6: Element Tracker
    # ----------------------------------------------------------------
    et = wb.create_sheet("6. Element Tracker")
    et.sheet_view.showGridLines = False
    col_widths(et, {"A": 2, "B": 9, "C": 12, "D": 16, "E": 40, "F": 12, "G": 16, "H": 16, "I": 8, "J": 22})
    et["B2"] = "Protection Element Tracker"; et["B2"].font = f_title
    for j, txt in zip("BCDEFGHIJ", ["ANSI", "ABB Function", "IEC 61850 ID", "Description",
                                    "Enabled? (Y/N)", "Last tested (date)",
                                    "Test this visit? (Y/N)", "Result", "Notes"]):
        hdr_cell(et, f"{j}5", txt)
    el_data = [
        ("87G", "GENPDIF", "GENPDIF", "Generator differential"),
        ("87T", "T2WPDIF", "T2WPDIF", "Unit transformer differential"),
        ("87N", "REFPDIF", "REFPDIF", "Restricted earth fault (low-Z)"),
        ("64S", "STTIPHIZ", "STTIPHIZ", "100% stator earth fault (20 Hz inj.)"),
        ("64R", "ROTIPHIZ", "ROTIPHIZ", "Rotor earth fault (injection)"),
        ("40", "LEXPDIS", "LEXPDIS", "Loss of excitation (impedance)"),
        ("78", "PSPPPAM", "PSPPPAM", "Pole slip (lens)"),
        ("51V", "VRPVOC", "VRPVOC", "Volt-restrained OC (stator backup)"),
        ("50/51", "OC4PTOC", "OC4PTOC", "Phase overcurrent backup"),
        ("50N/51N", "EF4PTOC", "EF4PTOC", "Residual overcurrent backup"),
        ("46", "NS4PTOC", "NS4PTOC", "Negative-sequence overcurrent"),
        ("49S", "GSPTTR", "GSPTTR", "Stator thermal overload (I²t)"),
        ("49R", "GRPTTR", "GRPTTR", "Rotor thermal overload (I²t)"),
        ("27", "UV2PTUV", "UV2PTUV", "Undervoltage (2-step)"),
        ("59", "OV2PTOV", "OV2PTOV", "Overvoltage (2-step)"),
        ("59N", "ROV2PTOV", "ROV2PTOV", "Residual overvoltage"),
        ("24", "OEXPVPH", "OEXPVPH", "Overexcitation (V/Hz)"),
        ("32", "GOPPDOP", "GOPPDOP", "Directional overpower"),
        ("37", "GUPPDUP", "GUPPDUP", "Directional underpower"),
        ("50BF", "CCRBRF", "CCRBRF", "Breaker failure (3-phase)"),
        ("52PD", "CCPDSC", "CCPDSC", "CB pole discrepancy"),
        ("25", "SESRSYN", "SESRSYN", "Synchronism / energising check"),
        ("50AE", "AEGPVOC", "AEGPVOC", "Accidental energising"),
        ("81U", "SAPTUF", "SAPTUF", "Underfrequency"),
        ("81O", "SAPTOF", "SAPTOF", "Overfrequency"),
        ("81R", "SAPFRC", "SAPFRC", "ROCOF"),
        ("60", "FUFSPVC", "FUFSPVC", "Fuse / VT failure supervision"),
        ("CCS", "CCS", "CCS", "Current circuit supervision"),
    ]
    for i, (ansi, abb, iecid, desc) in enumerate(el_data):
        r = 6 + i
        c = et[f"B{r}"]; c.value = ansi; c.font = f_label; c.fill = fill_accent; c.alignment = ctr; c.border = border
        for j, v in zip("CD", [abb, iecid]):
            cc = et[f"{j}{r}"]; cc.value = v; cc.font = f_body; cc.alignment = ctr; cc.border = border
        cc = et[f"E{r}"]; cc.value = desc; cc.font = f_body; cc.alignment = left; cc.border = border
        for j in "FGHIJ":
            inputcell(et, f"{j}{r}")
    et.freeze_panes = "B6"

    # ----------------------------------------------------------------
    # SHEET 7: Element Results
    # ----------------------------------------------------------------
    er = wb.create_sheet("7. Element Results")
    er.sheet_view.showGridLines = False
    col_widths(er, {"A": 2, "B": 5, "C": 9, "D": 22, "E": 22, "F": 12, "G": 14,
                    "H": 11, "I": 13, "J": 14, "K": 11, "L": 8, "M": 22})
    er["B2"] = "Per-Element Test Results"; er["B2"].font = f_title
    for j, txt in zip("BCDEFGHIJKLM", ["#", "ANSI", "ABB Function", "Setting tested",
                                       "Pickup setting", "Pickup measured", "Pickup err %",
                                       "Time setting (s)", "Time measured (s)", "Time err %",
                                       "Result", "Notes"]):
        hdr_cell(er, f"{j}5", txt)
    for i in range(16):
        r = 6 + i
        c = er[f"B{r}"]; c.value = i + 1; c.font = f_body; c.alignment = ctr; c.border = border
        for j in "CDEFGIJ":
            inputcell(er, f"{j}{r}"); er[f"{j}{r}"].alignment = ctr if j in "FGIJ" else left
        # pickup err
        hc = er[f"H{r}"]
        hc.value = f'=IFERROR(IF(OR(F{r}="",G{r}=""),"",(G{r}-F{r})/F{r}*100),"")'
        hc.font = f_body; hc.alignment = ctr; hc.border = border; hc.fill = PatternFill("solid", fgColor="F2F2F2")
        # time err
        kc = er[f"K{r}"]
        kc.value = f'=IFERROR(IF(OR(I{r}="",J{r}=""),"",(J{r}-I{r})/I{r}*100),"")'
        kc.font = f_body; kc.alignment = ctr; kc.border = border; kc.fill = PatternFill("solid", fgColor="F2F2F2")
        # result
        lc = er[f"L{r}"]
        lc.value = (f'=IF(AND(NOT(ISNUMBER(H{r})),NOT(ISNUMBER(K{r}))),"",'
                    f'IF(OR(AND(ISNUMBER(H{r}),ABS(N(H{r}))>TolPickup_pct),'
                    f'AND(ISNUMBER(K{r}),ABS(N(K{r}))>TolTime_pct)),"FAIL","PASS"))')
        lc.font = f_body; lc.alignment = ctr; lc.border = border
        inputcell(er, f"M{r}")
    er.conditional_formatting.add("L6:L21", PASS_RULE())
    er.conditional_formatting.add("L6:L21", FAIL_RULE())
    er.freeze_panes = "B6"

    # ----------------------------------------------------------------
    # SHEET 8: Post-Test
    # ----------------------------------------------------------------
    po = wb.create_sheet("8. Post-Test")
    po.sheet_view.showGridLines = False
    col_widths(po, {"A": 2, "B": 6, "C": 60, "D": 12, "E": 26, "F": 10})
    po["B2"] = "Post-Test Checklist"; po["B2"].font = f_title
    n = 1
    r = checklist_block(po, 4, "8.1 Settings re-comparison (as-left vs. master)", [
        "PCM600 upload; saved as asleft_[date]_[tag].pcm.",
        "Compared as-left vs. master settings.",
        "Zero unauthorised differences remain.",
        "Temporarily disabled / modified functions restored.",
    ], n); n += 4
    r = checklist_block(po, r + 1, "8.2 Clear alarms / reset indications", [
        "Test-mode alarms cleared via LHMI > Clear.",
        "LHMI alarm LEDs and latched indications reset.",
        "Internal events / disturbance buffers reviewed and archived.",
    ], n); n += 3
    r = checklist_block(po, r + 1, "8.3 Time synchronisation", [
        "Re-read IED date/time.",
        "If drift outside tolerance, corrected via configured time source.",
        "TIMESYNCHERROR inactive; sync healthy.",
        "Final drift within tolerance.",
    ], n); n += 4
    r = checklist_block(po, r + 1, "8.4 Plant restoration (REG670-specific)", [
        "Disconnected all test leads from IED terminals.",
        "Replaced all test block links / reinstated VT MCBs.",
        "Removed CT shorting links; verified continuity to source.",
        "Restored trip and lockout output links / re-enabled trips.",
        'Self-supervision "Normal" green LED re-lit and steady.',
        "Cleared any test-mode alarms via LHMI > Clear.",
        "INTERNAL FAIL contact de-energised (healthy).",
        "Active setting group restored to correct group (not test group).",
        "64S/64R injection systems (REX060) re-enabled if blocked.",
        "No targets / alarm LEDs remain active on IED or panel.",
        "SCADA / DCS analogue values match IED metering live.",
        "Re-energised VTs; sensible secondary voltages on IED.",
        "Operations / control room notified work complete.",
        "Permit returned and isolations removed (LOTO removed).",
    ], n); n += 14
    r = checklist_block(po, r + 1, "8.5 Final self-supervision check", [
        "Diagnostics > IED status > General: no FAIL / WARNING.",
        '"Normal" green LED lit and steady.',
        "IED in service / not in test mode per plant requirements.",
        "Final screenshots captured (status / Measurements / LEDs).",
    ], n); n += 4

    # ----------------------------------------------------------------
    # SHEET 9: Files & Sign-Off
    # ----------------------------------------------------------------
    so = wb.create_sheet("9. Files & Sign-Off")
    so.sheet_view.showGridLines = False
    col_widths(so, {"A": 2, "B": 34, "C": 26, "D": 14, "E": 12})
    so["B2"] = "Findings Log & Sign-Off"; so["B2"].font = f_title

    section_title(so, "B4", "Findings, Defects & Recommendations (detailed log)")
    for j, txt in zip("BCDE", ["#", "Finding / Defect", "Severity", "Closed?"]):
        hdr_cell(so, f"{j}5", txt)
    for i in range(8):
        r = 6 + i
        c = so[f"B{r}"]; c.value = i + 1; c.font = f_body; c.alignment = ctr; c.border = border
        for j in "CDE":
            inputcell(so, f"{j}{r}")

    section_title(so, "B16", "Test Files Index")
    files = ["Events (EVE)", "Disturbance records (DRPRDRE)", "Internal events list",
             "As-found PCM600 project", "As-left PCM600 project", "Omicron test file (.occ)",
             "Screenshots"]
    for j, txt in zip("BCD", ["File type", "File name", "Location"]):
        hdr_cell(so, f"{j}17", txt)
    for i, fl in enumerate(files):
        r = 18 + i
        c = so[f"B{r}"]; c.value = fl; c.font = f_label; c.fill = fill_accent; c.alignment = left; c.border = border
        inputcell(so, f"C{r}"); inputcell(so, f"D{r}")

    so_signrow = 18 + len(files) + 1
    section_title(so, f"B{so_signrow}", "Sign-Off")
    sign = ["Tester name", "Tester signature", "Date", "Witness name", "Witness signature",
            "Date", "Authorising engineer (acceptance)", "Signature", "Date"]
    for i, lab in enumerate(sign):
        r = so_signrow + 1 + i
        label(so, f"B{r}", lab); inputcell(so, f"C{r}")
    orow = so_signrow + 1 + len(sign)
    label(so, f"B{orow}", "Overall test result (from Cover sheet)")
    c = so[f"C{orow}"]; c.value = f"={overall_ref}"; c.font = f_label; c.alignment = left; c.border = border

    # ----------------------------------------------------------------
    # Defined names
    # ----------------------------------------------------------------
    for name, coord in {**rmap, **tolmap}.items():
        wb.defined_names.add(DefinedName(name, attr_text=f"'1. Cover'!${coord[0]}${coord[1:]}"))

    # tab colours
    for ws in wb.worksheets:
        ws.sheet_properties.tabColor = BRAND

    out = os.path.join(HERE, "REG670_Test_Results_Template.xlsx")
    wb.save(out)
    return out


if __name__ == "__main__":
    docx_path = build_docx()
    print("Wrote:", docx_path)
    xlsx_path = build_xlsx()
    print("Wrote:", xlsx_path)
